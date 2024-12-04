import os
from pathlib import Path
from docx import Document
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SimpleNodeParser

def process_docx_folder(input_folder, output_folder):
    """处理文件夹内的所有DOCX文件，分割并保存"""
    # 如果输出文件夹不存在，则创建
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    for root, dirs, files in os.walk(input_folder):
        # 保持原有的文件夹结构
        relative_path = os.path.relpath(root, input_folder)
        output_subfolder = os.path.join(output_folder, relative_path)
        Path(output_subfolder).mkdir(parents=True, exist_ok=True)

        for file in files:
            # 忽略以 ~$ 开头的临时文件
            if file.startswith("~$"):
                continue

            if file.lower().endswith('.docx'):
                input_path = os.path.join(root, file)
                output_path = os.path.join(output_subfolder, file)

                # 读取原始DOCX文件
                try:
                    doc = Document(input_path)
                    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])  # 获取非空段落的文本

                    # 使用LlamaIndex进行文本分割
                    sentences = split_text_using_llamaindex(text, max_paragraph_length=400)

                    # 将分割后的文本保存到新的DOCX文件
                    save_text_as_docx(sentences, output_path)
                    print(f"Processed: {input_path} -> {output_path}")
                except Exception as e:
                    print(f"Error processing {input_path}: {e}")


def split_text_using_llamaindex(text, max_paragraph_length=500):
    """使用LlamaIndex对文本进行逻辑分割，每个段落长度小于max_paragraph_length"""
    try:
        # 创建 LlamaDocument
        document = LlamaDocument(text=text)  # 使用命名参数
    except TypeError as e:
        print(f"Error: {e}")  # 捕获类型错误并输出详细信息
        return []

    # 使用 SimpleNodeParser 解析文档
    parser = SimpleNodeParser()
    nodes = parser.get_nodes_from_documents([document])

    sentences = []
    current_paragraph = []

    # 将每个节点的文本按最大字数限制分割
    for node in nodes:
        node_text = node.text.strip()

        # 如果当前段落长度加上当前节点超出限制，就保存当前段落并开始一个新段落
        if len("\n".join(current_paragraph) + "\n" + node_text) > max_paragraph_length:
            sentences.append("\n".join(current_paragraph))
            current_paragraph = [node_text]
        else:
            current_paragraph.append(node_text)

    # 最后添加一个段落
    if current_paragraph:
        sentences.append("\n".join(current_paragraph))

    return sentences


def save_text_as_docx(sentences, output_path):
    """将分割后的段落保存为新的DOCX文件"""
    doc = Document()
    for sentence in sentences:
        doc.add_paragraph(sentence)
        doc.add_paragraph("###")  # 添加自定义的分隔符
    doc.save(output_path)

if __name__ == "__main__":
    input_folder = r"D:/Desktop/1"  # 请替换为你的DOCX文件夹路径
    output_folder = r"D:/Desktop/segments"  # 请替换为你希望保存分割后文件的路径
    process_docx_folder(input_folder, output_folder)
