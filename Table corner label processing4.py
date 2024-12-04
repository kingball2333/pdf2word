import os
import glob
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def table_to_text(table):
    """
    将.docx文档中的表格转换为纯文本格式。

    参数:
    - table: 一个 docx 表格对象。

    返回:
    - 一个字符串，其中每一行的单元格用 '|' 连接，每行之间用换行符分隔。
    """
    text = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        text.append('|'.join(cells))
    return '\n'.join(text)


def set_font(para, font_name='微软雅黑', font_size=11):
    """
    设置段落中所有文本的字体样式。

    参数:
    - para: 一个 docx 段落对象。
    - font_name: 字体名称 。
    - font_size: 字体大小 。
    """
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)  # 设置英文字体


def process_subscript_superscript(doc, font_name, font_size):
    """
    处理文档中的上下标，将下标用 <sub> 标签包围，将上标用 <sup> 标签包围。
    同时设置字体样式。

    参数:
    - doc: 一个 docx Document 对象。
    - font_name: 字体名称。
    - font_size: 字体大小。
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.subscript:
                run.text = f'<sub>{run.text}</sub>'
                run.font.name = font_name
                run.font.size = Pt(font_size)
            elif run.font.superscript:
                run.text = f'<sup>{run.text}</sup>'
                run.font.name = font_name
                run.font.size = Pt(font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.subscript:
                            run.text = f'<sub>{run.text}</sub>'
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                        elif run.font.superscript:
                            run.text = f'<sup>{run.text}</sup>'
                            run.font.name = font_name
                            run.font.size = Pt(font_size)


def replace_tables_with_text(doc):
    """
    将 .docx 文档中的所有表格替换为其文本表示形式。

    参数:
    - doc: 一个 docx Document 对象。
    """
    tables = list(doc.tables)

    for table in tables:
        table_text = table_to_text(table)

        p = doc.add_paragraph()
        p.add_run(table_text)
        set_font(p, '微软雅黑', 11)

        parent = table._element.getparent()
        parent.insert(parent.index(table._element), p._element)

        blank_p = OxmlElement('w:p')
        parent.insert(parent.index(p._element) + 1, blank_p)

        parent.remove(table._element)


def save_document(doc, output_path):
    """
    保存修改后的文档到指定路径。

    参数:
    - doc: 一个 docx Document 对象。
    - output_path: 保存文档的路径。
    """
    doc.save(output_path)


def process_documents(input_folder, output_folder, font_name='微软雅黑', font_size=11):
    """
    处理指定文件夹中的所有 .docx 文档，处理上下标，替换表格为文本，并保存到输出文件夹。

    参数:
    - input_folder: 输入文件夹的路径，包含要处理的 .docx 文件。
    - output_folder: 输出文件夹的路径，保存处理后的文档。
    - font_name: 字体名称。
    - font_size: 字体大小。
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for input_file in glob.glob(os.path.join(input_folder, "*.docx")):
        doc = Document(input_file)

        process_subscript_superscript(doc, font_name, font_size)
        replace_tables_with_text(doc)

        base_name = os.path.basename(input_file)
        output_file = os.path.join(output_folder, base_name)

        save_document(doc, output_file)
        print(f"已处理: {output_file}")


# 定义输入和输出文件夹路径
input_folder = r"D:\Desktop\传染病 - 结束"
output_folder = r"D:\Desktop\传染病_处理"

# 处理输入文件夹中的所有文档
process_documents(input_folder, output_folder)