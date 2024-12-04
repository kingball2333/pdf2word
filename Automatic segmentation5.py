import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
import os
import glob
import shutil

# 标题匹配模式
pattern_1 = re.compile(r'^\s*\d+\s*[\u4e00-\u9fa5\uf900-\ufa2d]+$')  # 一级标题
pattern_2 = re.compile(r'^\s*\d+\.\d(?!\.)(?!\d)\s*[\u4e00-\u9fa5\uf900-\ufa2d]*')  # 二级标题
pattern_3 = re.compile(r'^\s*\d+\.\d\.\d(?!\.)(?!\d)\s*[\u4e00-\u9fa5\uf900-\ufa2d]?')  # 三级标题
pattern_4 = re.compile(r'^\s*附录\s?[A-Z]$')  # 附录
pattern_cn = re.compile(r'^\s*第[\u4e00-\u9fa5]+章\s*')  # 中文标题


def find_heading(path):
    """查找文档中的所有标题并验证其顺序"""
    doc = Document(path)
    heading = {}
    cnt = 0

    # 查找标题并记录位置
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if pattern_1.match(text) or pattern_2.match(text) or pattern_3.match(text) or pattern_4.match(text):
            heading[text] = cnt
        cnt += 1

    # 提取标题的第一个数字用于检查顺序
    title_first = [int(re.match(r'^(\d+)', heading).group(1)) for heading in heading.keys() if
                   re.match(r'^(\d+)', heading)]
    # 检查标题顺序
    for i in range(len(title_first) - 1):
        if title_first[i] > title_first[i + 1]:
            return f"标题顺序错误: '{list(heading.keys())[i - 1], list(heading.keys())[i], list(heading.keys())[i + 1]}'"
    return heading


def find_main_heading_cn(path):
    """查找主要标题"""
    doc = Document(path)
    main_heading = {}
    cnt = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if pattern_cn.match(text):
            main_heading[text] = cnt
        cnt += 1
    return main_heading


def words_main_count_cn(path):
    """计算文档中各主要标题部分的字数"""
    doc = Document(path)
    paragraphs = doc.paragraphs
    word_counts = [len(paragraph.text) for paragraph in paragraphs]

    headings = find_main_heading_cn(path)
    titles = list(headings.keys())
    pages = list(headings.values())

    results = {}
    if titles:
        first_heading_start = pages[0]
        introduction_word_count = sum(word_counts[:first_heading_start])
        results["0 介绍"] = [0, first_heading_start - 1, introduction_word_count]

        for i in range(len(pages)):
            start_index = pages[i]
            end_index = pages[i + 1] if i + 1 < len(pages) else len(word_counts)
            section_word_count = sum(word_counts[start_index:end_index])
            results[titles[i]] = [start_index, end_index - 1, section_word_count]
    return results


def find_main_heading(path):
    """查找主要标题"""
    doc = Document(path)
    main_heading = {}
    cnt = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if pattern_1.match(text) or pattern_4.match(text):
            main_heading[text] = cnt
        cnt += 1
    return main_heading


def words_main_count(path):
    """计算文档中各主要标题部分的字数"""
    doc = Document(path)
    paragraphs = doc.paragraphs
    word_counts = [len(paragraph.text) for paragraph in paragraphs]

    headings = find_main_heading(path)
    titles = list(headings.keys())
    pages = list(headings.values())

    results = {}
    if titles:
        first_heading_start = pages[0]
        introduction_word_count = sum(word_counts[:first_heading_start])
        results["0 介绍"] = [0, first_heading_start - 1, introduction_word_count]

        for i in range(len(pages)):
            start_index = pages[i]
            end_index = pages[i + 1] if i + 1 < len(pages) else len(word_counts)
            section_word_count = sum(word_counts[start_index:end_index])
            results[titles[i]] = [start_index, end_index - 1, section_word_count]
    return results


def find_second_heading(path):
    """查找二级标题"""
    doc = Document(path)
    second_heading = {}
    cnt = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if pattern_2.match(text):
            second_heading[text] = cnt
        cnt += 1
    return second_heading


def words_second_count(path):
    """计算文档中各二级标题部分的字数"""
    doc = Document(path)
    paragraphs = doc.paragraphs
    word_counts = [len(paragraph.text) for paragraph in paragraphs]

    headings = find_second_heading(path)
    titles = list(headings.keys())
    pages = list(headings.values())

    results = {}
    if titles:
        for i in range(len(pages) - 1):
            start_index = pages[i]
            end_index = pages[i + 1]
            section_word_count = sum(word_counts[start_index:end_index])
            results[titles[i]] = [start_index, end_index - 1, section_word_count]

        # 处理最后一个二级标题
        last_title = titles[-1]
        last_start_index = pages[-1]

        # 查找下一个一级标题或文档结尾的位置
        end_index = len(paragraphs)
        for j in range(last_start_index, len(paragraphs)):
            text = paragraphs[j].text.strip()
            if pattern_1.match(text) or pattern_4.match(text):
                end_index = j
                break

        section_word_count = sum(word_counts[last_start_index:end_index])
        results[last_title] = [last_start_index, end_index - 1, section_word_count]

    return results


def find_third_heading(path):
    """查找三级标题"""
    doc = Document(path)
    third_heading = {}
    cnt = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if pattern_3.match(text):
            third_heading[text] = cnt
        cnt += 1
    return third_heading


def words_third_count(path):
    """计算文档中各三级标题部分的字数"""
    doc = Document(path)
    paragraphs = doc.paragraphs
    word_counts = [len(paragraph.text) for paragraph in paragraphs]
    headings = find_third_heading(path)
    all_headings = find_heading(path)
    titles = list(all_headings.keys())
    pages = list(all_headings.values())

    results = {}
    if titles:
        for i in range(len(pages) - 1):
            start_index = pages[i]
            end_index = pages[i + 1]
            section_word_count = sum(word_counts[start_index:end_index])
            results[titles[i]] = [start_index, end_index - 1, section_word_count]
    return {k: v for k, v in results.items() if k in headings}


def find_struct(path):
    """构建标题的结构"""
    headings = find_heading(path)
    if isinstance(headings, str):
        return headings  # 返回标题顺序错误信息

    main_results = words_main_count(path)
    second_results = words_second_count(path)
    third_results = words_third_count(path)
    # 合并结果
    results = {**main_results, **second_results, **third_results}
    if len(results) == 0:
        results = headings
    struct = []
    stack = [struct]  # 初始化栈

    def create_entry(title, content):
        return {
            'title': title,
            'content': content,
            'children': []
        }

    processed_titles = set()

    # 添加“介绍”到结构中
    if '0 介绍' in results:
        struct.append(create_entry('0 介绍', results['0 介绍']))
    # 处理标题
    for heading in headings.keys():
        if heading in results:
            match = re.search(r'\d+(?:\.\d+)*', heading)
            if match:
                level_str = match.group(0)
                level = len(level_str.split('.')) - 1
                content = results[heading]
                entry = create_entry(heading, content)

                while len(stack) > level + 1:
                    stack.pop()

                if heading not in processed_titles:
                    stack[-1].append(entry)
                    processed_titles.add(heading)
                    stack.append(entry['children'])

    # 处理附录
    for heading in headings.keys():
        if heading not in processed_titles and re.search(r'^附录', heading):
            content = results.get(heading, "")
            struct.append(create_entry(heading, content))
            processed_titles.add(heading)

    # 清理空的条目
    def clean_up(lst):
        i = 0
        while i < len(lst):
            if isinstance(lst[i], dict) and not lst[i]['title'] and not lst[i]['children']:
                lst.pop(i)
            else:
                if 'children' in lst[i]:
                    clean_up(lst[i]['children'])
                i += 1

    clean_up(struct)
    return struct


def print_struct(result, level=0):
    """递归打印结构化数据"""
    indent = ' ' * (level * 2)

    if isinstance(result, dict):
        if result['title'] or result['content']:
            print(f"{indent}{result['title']}  {result['content']}")
        for child in result.get('children', []):
            print_struct(child, level + 1)
    elif isinstance(result, list):
        for item in result:
            print_struct(item, level)


def modify_font(paragraph, font_name, font_size, font_color):
    """修改段落的字体属性"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*font_color)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)


def file_process(path, save_path):
    """处理单个文档，修改标题并保存"""
    doc = Document(path)
    titles = set()
    if len(find_main_heading(path)) != 0:
        try:
            struct = find_struct(path)
            if isinstance(struct, str):
                return struct  # 返回错误信息
            # 收集标题
            for item in struct:
                if len(item['content']) > 2 and item['content'][2] <= 5000:
                    titles.add(item['title'])
                else:
                    if item['children']:
                        for child in item['children']:
                            if len(child['content']) > 2 and child['content'][2] <= 5000:
                                titles.add(child['title'])
                            else:
                                if child['children']:
                                    for childed in child['children']:
                                        if len(childed['content']) > 2:
                                            titles.add(childed['title'])
                                else:
                                    titles.add(child['title'])
                    else:
                        titles.add(item['title'])
            # 修改文档中的标题
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text in titles:
                    for run in paragraph.runs:
                        if text in run.text:
                            run.text = '### ' + text
                            modify_font(paragraph, font_name='微软雅黑', font_size=11, font_color=(0, 0, 0))
                            break

            doc.save(save_path)
        except Exception as e:
            return str(e)
    else:
        heading = find_main_heading_cn(path)
        headings = list(heading.keys())
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text in headings:
                for run in paragraph.runs:
                    if text in run.text:
                        run.text = '### ' + text
                        modify_font(paragraph, font_name='微软雅黑', font_size=11, font_color=(0, 0, 0))
                        break

        doc.save(save_path)


def process_documents(input_folder, output_folder, error_folder):
    """处理指定文件夹中的所有 .docx 文档，根据处理结果保存到不同文件夹"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    if not os.path.exists(error_folder):
        os.makedirs(error_folder)

    for input_file in glob.glob(os.path.join(input_folder, "*.docx")):
        base_name = os.path.basename(input_file)
        output_file = os.path.join(output_folder, base_name)
        error_file = os.path.join(error_folder, base_name)

        result = file_process(input_file, output_file)
        if isinstance(result, str):  # 处理失败
            shutil.move(input_file, error_file)  # 移动到错误文件夹
            print(f"处理失败: {input_file}，错误信息: {result}")
        else:  # 处理成功
            print(f"成功处理: {output_file}")


# 输入、输出和错误文件夹路径
input_folder = r"D:\Desktop\职业卫生_结束"
output_folder = r"D:\Desktop\结束0_分段"
error_folder = r"D:\Desktop\结束0_出错"
# process_documents(input_folder,output_folder,error_folder)


'''纠错环节'''


# def process_folder(folder_path):
#     """处理指定文件夹中的所有 .docx 文件"""
#     for file_path in glob.glob(os.path.join(folder_path, "*.docx")):
#         print(f"处理文件: {file_path}")
#         result = find_struct(file_path)
#         if isinstance(result, str):
#             print(f"错误: {result}")


# folder_path = r"D:\Desktop\职业卫生_结束"
# process_folder(folder_path)

# file_path = r"D:\Desktop\《口腔智慧监管》\消毒管理办法.docx"
# result = find_struct(file_path)
# if isinstance(result, str):
#     print(f"错误: {result}")
# else:
#     print_struct(result)
